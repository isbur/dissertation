from docx import *
from docx.shared import Inches
from glob import glob
from natsort import natsorted
import os 
from os.path import join
import re
import sys


# Preparations
# Get Postfix
if len(sys.argv)==2:
    Postfix = sys.argv[1]
else:
    Postfix = "A"
    print("Wrong number of arguments")
    print("Continue with A postfix")
# Get current file directory name
dirname = os.path.dirname(__file__)

# Parse labels
wd = Document( os.path.join(dirname, "../common/Anketa.docx") )
questionNumber = 0
Labels = []
for paragraph in wd.paragraphs:
    if paragraph.text.find("Другое") != -1:
        continue
    
    if paragraph.text.find("Вопрос") != -1:
        Labels.append([paragraph.text])
        questionNumber += 1
        continue
    
    #if paragraph.text.find("Сколько") != -1:
        #column = [cell.value for cell in [column for column in ws.iter_cols(min_col=questionNumber, max_col=questionNumber)][0]][1:]
        #Labels[len(Labels) - 1] = Labels[len(Labels) - 1] + list(np.unique(column))[2:]
        #continue
    
    if paragraph.style.name == 'List Paragraph':
        Labels[len(Labels) - 1].append(paragraph.text)
        continue


wd = Document()
section = wd.sections[0]
chiselko = shared.Inches(0.5)
section.left_margin, section.right_margin, section.top_margin, section.bottom_margin = chiselko, chiselko, chiselko, chiselko

fileList = glob(join(dirname, "./*" + Postfix + ".png"))
fileList = natsorted(fileList)
#for i, filename in enumerate(fileList):
    #wd.add_paragraph( Labels[i][0] )
    #wd.add_picture( filename )

#B = []
for i, label in enumerate(Labels):
    wd.add_paragraph( label[0] )
    A = [x for x in fileList if re.search(' '+str(i+1)+'((\..?.?)|\B)'+Postfix+'\.png', x)]
    for j, filename in enumerate(A):
        if len(A)>1:
            wd.add_paragraph(str(j+1) + '. ' + label[j+1])
        wd.add_picture( filename, height = Inches(4.5) )   
        
    #print(i)
    #A.append([x for x in fileList if re.search('. '+str(i+1)+'.?.?'+Postfix+'\.png', x) or print('x:\t',x) or print(re.search('. '+str(i+1)+'.?.?'+Postfix+'\.png', x))])
    #B.append(A)
    #print('\n', '\n'.join(list(map(str, B[i]))))
    
    #print('\n\n\n')

wd.save( os.path.join(dirname, 'summary' + Postfix + '.docx') )